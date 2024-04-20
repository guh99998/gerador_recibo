from docx import Document
from docx.shared import Pt
from modules.format_documento import formatar_doc
from modules.format_data import formatar_data
from modules.format_numero_extenso import formatar_numero_extenso
from constants import *

class Recibo:
    def __init__(self, pagador, recebedor, nome, valor, referencia, data, cidade, estado):
        '''
        Construtor da classe Recibo, exigindo parâmetros para a criação de um objeto do tipo Recibo.
        '''
        self.document = Document()
        self._pagador = pagador
        self._recebedor = recebedor
        self._nome = nome
        self._valor = valor
        self._referencia = referencia
        self._data = data
        self._cidade = cidade
        self._estado = estado

    def add_titulo(self):
        '''
        Adiciona um título no documento word gerado, o título recebe um alinhamento para ficar centralizado no documento criado.
        '''
        titulo = self.document.add_heading(TITULO, level=0)
        titulo.alignment = ALINHAMENTO_CENTRAL

    def add_valor(self):
        '''
        Adiciona um valor no documento word gerado, esse valor é recebido pelo próprio usuário e é customizado com um alinhamento
        para a direita e uma fonte padrão de valor. Além disso, também é adicionado um espaçamento antes do local onde esse valor fica
        e um espaçamento depois de onde esse valor fica.
        '''
        valor = self.document.add_paragraph()
        run_valor = valor.add_run(f"R${self._valor:.2f}")
        run_valor.font.size = VALOR_FONT_SIZE
        valor.alignment = ALINHAMENTO_DIREITA
        valor.paragraph_format.space_before = ESPACAMENTO_24
        valor.paragraph_format.space_after = ESPACAMENTO_72

    def add_corpo(self):
        '''
        Adiciona o corpo do recibo, com um alinhamento justificado e passa informações do cliente que está efetuado o pagamento (nome e documento),
        passa o valor que foi colocado no início do recibo, esse valor também é passado por extenso e após isso é informada a referência do pagamento.
        '''
        corpo = self.document.add_paragraph('Recebi de "')
        corpo.alignment = ALINHAMENTO_JUSTIFICADO
        run_corpo = corpo.add_run({self._pagador._nome.upper()})
        run_corpo.bold = True
        corpo.add_run(
            f'", inscrito no {self._pagador.tipo_pessoa()} n° {formatar_doc(self._pagador._documento)}, a importância de R${self._valor:.2f} ({formatar_numero_extenso(self._valor)}), referente ao {self._referencia}.'
        )

    def add_data_assinatura(self):
        '''
        Adiciona a cidade e estado do local da emissão do recibo, a data é formatada para ser escrita por extenso e fica centralizada no documento
        criado. Abaixo do local, é adicionada a linha de assinatura para ser utilizada no documento impresso.
        '''
        data = self.document.add_paragraph()
        data.alignment = ALINHAMENTO_CENTRAL
        data_run = data.add_run(
            f"{self._cidade} - {self._estado.upper()}, {formatar_data(self._data)}."
        )
        data.paragraph_format.space_after = ESPACAMENTO_72
        data.paragraph_format.space_before = ESPACAMENTO_50

        assinatura = self.document.add_paragraph()
        assinatura.add_run("\n\n")
        assinatura.add_run(
            "__________________________________________________________________________"
        )
        assinatura.alignment = ALINHAMENTO_CENTRAL

    def add_recebedor(self):
        '''
        Adiciona o cliente que vai receber o valor do recibo e o documento, o valor fica alinhado de forma centralizada
        e o documento é formatado para obter o tipo e passar o valor.
        '''
        nome_recebedor = self.document.add_paragraph()
        nome_recebedor.alignment = ALINHAMENTO_CENTRAL
        nome_recebedor_run = nome_recebedor.add_run(f"{self._recebedor.nome_upper()}")
        documento_recebedor = self.document.add_paragraph()
        documento_recebedor_run = documento_recebedor.add_run(
            f"{self._recebedor.tipo_pessoa()}: {formatar_doc(self._recebedor._documento)}"
        )
        documento_recebedor.alignment = ALINHAMENTO_CENTRAL

    def salvar_recibo(self):
        '''
        Salva o recibo criado.
        '''
        self.document.save(f"{self._nome}.docx")

    def criar_recibo(self):
        '''
        Junta todas as funções criadas para gerar o recibo.
        '''
        self.add_titulo()
        self.add_valor()
        self.add_corpo()
        self.add_data_assinatura()
        self.add_recebedor()